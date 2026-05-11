"""
================================================================================
 methodology_parser.py
 Universal Document Ingestion & Semantic Extraction Engine
================================================================================
"""

import argparse
import json
import logging
import os
import re
import urllib.request
import urllib.error
from pathlib import Path
from typing import Optional, Dict, Any

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger("semantic_engine")

# Optional Imports (Graceful Degradation)
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    import docx
except ImportError:
    docx = None

try:
    import polars as pl
except ImportError:
    pl = None


# ══════════════════════════════════════════════════════════════════════════════
#  INGESTION LAYER
# ══════════════════════════════════════════════════════════════════════════════

def extract_text_from_pdf(path: str) -> str:
    text = ""
    if pdfplumber:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        for row in table:
                            text += "| " + " | ".join(str(cell).replace('\n', ' ') if cell else "" for cell in row) + " |\n"
    elif fitz:
        doc = fitz.open(path)
        for page in doc:
            text += page.get_text("text") + "\n"
    else:
        raise ImportError("Install 'pdfplumber' or 'PyMuPDF' (fitz) to process PDF files.")
    return text

def extract_text_from_docx(path: str) -> str:
    if not docx:
        raise ImportError("Install 'python-docx' to process DOCX files.")
    doc = docx.Document(path)
    text = ""
    for para in doc.paragraphs:
        if para.text.strip():
            text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            text += "| " + " | ".join(cell.text.replace('\n', ' ') for cell in row.cells) + " |\n"
    return text

def extract_text_from_excel(path: str) -> str:
    if not pl:
        raise ImportError("Install 'polars' to process Excel files.")
    text = ""
    try:
        df = pl.read_excel(path)
        cols = df.columns
        text += "| " + " | ".join(cols) + " |\n"
        text += "| " + " | ".join(["---"] * len(cols)) + " |\n"
        for row in df.iter_rows():
            text += "| " + " | ".join(str(cell).replace('\n', ' ') if cell is not None else "" for cell in row) + " |\n"
    except Exception as e:
        logger.error(f"Failed to extract excel: {e}")
    return text

def parse_document(path: str) -> str:
    path_obj = Path(path)
    if not path_obj.exists():
        raise FileNotFoundError(f"File not found: {path}")
    
    ext = path_obj.suffix.lower()
    if ext == '.pdf':
        return extract_text_from_pdf(path)
    elif ext == '.docx':
        return extract_text_from_docx(path)
    elif ext in ['.xlsx', '.xls', '.csv']:
        if ext == '.csv' and pl:
            df = pl.read_csv(path)
            cols = df.columns
            text = "| " + " | ".join(cols) + " |\n"
            text += "| " + " | ".join(["---"] * len(cols)) + " |\n"
            for row in df.iter_rows():
                text += "| " + " | ".join(str(cell).replace('\n', ' ') if cell is not None else "" for cell in row) + " |\n"
            return text
        elif ext == '.csv':
            with open(path, 'r', encoding='utf-8') as f:
                return f.read()
        return extract_text_from_excel(path)
    elif ext in ['.md', '.txt', '.json']:
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported document format: {ext}")


# ══════════════════════════════════════════════════════════════════════════════
#  LLM EXTRACTION LAYER
# ══════════════════════════════════════════════════════════════════════════════

SYSTEM_PROMPT = """
You are an expert Data Architect. Your job is to read methodology documents, data dictionaries, and info-models provided as Markdown text, and extract the underlying data schema strictly as a JSON object.

Extract:
1. Column definitions (expected data type, descriptions).
2. Allowed values / business constraints for categorical columns.
3. Primary keys.
4. Foreign key relationships.

Output ONLY valid JSON matching this schema:
{
  "columns": {
    "column_name": {
      "expected_type": "string|int|float|boolean|date|datetime",
      "description": "Brief definition",
      "is_primary_key": true/false,
      "allowed_values": ["Value1", "Value2"] // only if categorical/enum
    }
  },
  "relationships": [
    {
      "source_col": "column_name",
      "target_table": "target_table_name",
      "target_col": "target_column_name"
    }
  ]
}
"""

def extract_schema_with_intern(text: str) -> Dict[str, Any]:
    try:
        import sys
        from pathlib import Path
        
        # Add root project to sys.path to allow imports if script run directly
        root_dir = str(Path(__file__).parent.parent)
        if root_dir not in sys.path:
            sys.path.insert(0, root_dir)
            
        from core.config import load
        from core.registry import InternRegistry
        
        cfg = load()
        registry = InternRegistry(cfg)
        intern = registry.get_intern("methodology_analyst")
        
        response_text = intern.run("", {"document_text": text})
        
        if not response_text or response_text.strip() == "{}":
            logger.warning("Intern returned empty response (possibly in fallback mode)")
            return {"columns": {}, "relationships": []}
            
        response_text = response_text.strip()
        response_text = re.sub(r"^```(?:json)?\s*|\s*```$", "", response_text.strip())
        return json.loads(response_text)
    except Exception as exc:
        logger.error(f"LLM extraction via intern failed: {exc}")
        return {"columns": {}, "relationships": []}


# ══════════════════════════════════════════════════════════════════════════════
#  CLI
# ══════════════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(description="Universal Document Ingestion & Semantic Extraction Engine")
    parser.add_argument("--doc", required=True, help="Path to methodology document (PDF, DOCX, Excel, MD, CSV)")
    parser.add_argument("--out", default="state/semantic_schema.json", help="Output JSON file path")
    args = parser.parse_args()

    logger.info(f"Parsing document: {args.doc}")
    try:
        raw_text = parse_document(args.doc)
        logger.info(f"Extracted {len(raw_text)} characters.")
    except Exception as e:
        logger.error(f"Document parsing failed: {e}")
        return

    logger.info("Extracting semantic schema via Intern Agent...")
    schema = extract_schema_with_intern(raw_text)
    
    out_path = Path(args.out)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(schema, f, indent=2)
    
    logger.info(f"Semantic schema saved to {out_path}")

if __name__ == "__main__":
    main()
